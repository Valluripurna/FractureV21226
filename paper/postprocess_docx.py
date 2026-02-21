from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

BASE_DIR = Path(__file__).parent
DOCX_PATH = str(BASE_DIR / 'fracturedetect_ai_paper.docx')

def set_section_single_column(section):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), '1')
    # keep spacing default

def insert_section_break_after(paragraph, two_columns=True, space_pts=18, continuous=True):
    p = paragraph._p
    # Ensure paragraph properties element exists
    pPr = p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    if continuous:
        type_el = OxmlElement('w:type')
        type_el.set(qn('w:val'), 'continuous')
        sectPr.append(type_el)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2' if two_columns else '1')
    cols.set(qn('w:space'), str(space_pts))
    sectPr.append(cols)
    pPr.append(sectPr)

def force_black_times(paragraph, size_pt=None, bold=None):
    for run in paragraph.runs:
        font = run.font
        font.name = 'Times New Roman'
        font.color.rgb = RGBColor(0, 0, 0)
        if size_pt is not None:
            font.size = Pt(size_pt)
        if bold is not None:
            font.bold = bold

def apply_references_format(paragraph):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(12)
    pf.left_indent = Inches(0.25)
    # Hanging indent: negative first line indent
    pf.first_line_indent = Inches(-0.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    force_black_times(paragraph, size_pt=9, bold=False)

def main():
    doc = Document(DOCX_PATH)

    # First section single column
    set_section_single_column(doc.sections[0])

    # Center Title and Author block, apply styles and insert continuous section break after last author email
    title_text = 'FractureDetect AI: A Deep Learning System for Automated Bone Fracture Detection in X-Ray Images'
    author_emails = {
        'purnavalluri03@gmail.com',
        'jaswanthsaisunkara1919@gmail.com',
        'amarpolisetti@gmail.com',
        'meganadhtata30@gmail.com',
    }

    # Style title
    for p in doc.paragraphs:
        if p.text.strip() == title_text:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            force_black_times(p, size_pt=24, bold=True)
            break

    author_names = [
        'Purna Chandra Rao Valluri',
        'Jaswanth Sai Sunkara',
        'Amarsai Polisetti',
        'Mohan Meganadh Tata',
    ]
    aff_lines = [
        'Department of Artificial Intelligence and Machine Learning',
        'NRI Institute of Technology',
        'Vijayawada, India',
    ]
    last_email_idx = None
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        # Emails
        if txt in author_emails:
            last_email_idx = i
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            force_black_times(p, size_pt=10, bold=False)
        # Names
        elif txt in author_names:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            force_black_times(p, size_pt=11, bold=False)
        # Affiliations
        elif txt in aff_lines:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            force_black_times(p, size_pt=10, bold=False)
            # Make affiliation italic (optional)
            for run in p.runs:
                run.font.italic = True

    # Abstract formatting: heading 10pt bold, content 9pt justified with left/right indent 0.25"
    abstract_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith('Abstract—'):
            abstract_idx = i
            for run in p.runs:
                text = run.text
                if text.strip().startswith('Abstract—'):
                    run.font.bold = True
                    run.font.size = Pt(10)
                else:
                    run.font.bold = False
                    run.font.size = Pt(9)
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.left_indent = Inches(0.25)
            pf.right_indent = Inches(0.25)
            break

    # Keywords formatting: heading bold 9pt, content 9pt italic
    keywords_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith('Keywords—'):
            keywords_idx = i
            for run in p.runs:
                text = run.text
                if text.strip().startswith('Keywords—'):
                    run.font.bold = True
                    run.font.size = Pt(9)
                else:
                    run.font.bold = False
                    run.font.italic = True
                    run.font.size = Pt(9)
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            break

    # Insert continuous two-column break after last author email (so all remaining content is two columns)
    if last_email_idx is not None:
        insert_section_break_after(doc.paragraphs[last_email_idx], two_columns=True, space_pts=18, continuous=True)

    # Apply body formatting to all paragraphs after the two-column break
    def apply_body_format(p):
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        force_black_times(p, size_pt=10, bold=False)

    if last_email_idx is not None:
        for p in doc.paragraphs[last_email_idx+1:]:
            # Avoid overriding captions/headings
            style_name = p.style.name if p.style else ''
            if style_name in ('Heading 1', 'Heading 2', 'Heading 3', 'Title', 'Caption'):
                continue
            apply_body_format(p)

    # Enforce black Times and sizes for headings/subheadings
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ''
        if style_name in ('Title', 'Heading 1', 'Heading 2', 'Heading 3'):
            size = 24 if style_name == 'Title' else 10
            force_black_times(p, size_pt=size, bold=(style_name != 'Title'))
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if style_name in ('Heading 1', 'Heading 2'):
                pf = p.paragraph_format
                pf.space_before = Pt(12)
                pf.space_after = Pt(6)

    # References formatting: find heading 'REFERENCES' and apply format to following paragraphs
    refs_start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == 'REFERENCES':
            refs_start = i + 1
            break
    if refs_start is not None:
        for p in doc.paragraphs[refs_start:]:
            # Stop if a new major heading appears (e.g., ACKNOWLEDGMENT)
            if p.style and p.style.name in ('Heading 1', 'Heading 2'):
                break
            apply_references_format(p)

    # Table formatting: 9pt center for all table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    force_black_times(para, size_pt=9, bold=False)

    # Figure captions: center 9pt
    for p in doc.paragraphs:
        if p.text.strip().lower().startswith('fig '):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            force_black_times(p, size_pt=9, bold=False)

    doc.save(DOCX_PATH)
    print('Post-processed DOCX: title/authors single-column, body two-columns (continuous), references formatted.')

if __name__ == '__main__':
    main()
